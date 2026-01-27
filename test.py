import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import os
import openai
import threading
from docx import Document  # For handling .docx files
from pypdf import PdfReader  # For handling .pdf files

# --- CONFIGURATION ---
openai.api_key = "KlPuOX3MvbmNVUYAGjWaEnZYOreipYZyRigY4oTXIYM"
openai.api_base = "https://api.poe.com/v1"
MODEL_NAME = "gemini-3-pro"


class DocGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Code Documenter")
        self.root.geometry("600x500")

        # --- UI ELEMENTS ---

        # 1. Code Folder Selection
        tk.Label(root, text="Step 1: Select Code Folder", font=("Arial", 10, "bold")).pack(pady=(10, 0))
        self.folder_path_var = tk.StringVar()
        frame1 = tk.Frame(root)
        frame1.pack(pady=5)
        tk.Entry(frame1, textvariable=self.folder_path_var, width=50).pack(side=tk.LEFT, padx=5)
        tk.Button(frame1, text="Browse", command=self.select_folder).pack(side=tk.LEFT)

        # 2. Template File Selection
        tk.Label(root, text="Step 2: Select Template (PDF or DOCX)", font=("Arial", 10, "bold")).pack(pady=(10, 0))
        self.template_path_var = tk.StringVar()
        frame2 = tk.Frame(root)
        frame2.pack(pady=5)
        tk.Entry(frame2, textvariable=self.template_path_var, width=50).pack(side=tk.LEFT, padx=5)
        tk.Button(frame2, text="Browse", command=self.select_template).pack(side=tk.LEFT)

        # 3. Generate Button
        self.generate_btn = tk.Button(root, text="GENERATE DOCUMENTATION", bg="#4CAF50", fg="white",
                                      font=("Arial", 12, "bold"), command=self.start_generation_thread)
        self.generate_btn.pack(pady=20)

        # 4. Logs
        tk.Label(root, text="Status Logs:", font=("Arial", 9)).pack(anchor="w", padx=20)
        self.log_area = scrolledtext.ScrolledText(root, height=10, width=70, state='disabled')
        self.log_area.pack(pady=5)

    # --- GUI ACTIONS ---

    def log(self, message):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def select_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path_var.set(folder)

    def select_template(self):
        file = filedialog.askopenfilename(filetypes=[("Documents", "*.pdf *.docx")])
        if file:
            self.template_path_var.set(file)

    def start_generation_thread(self):
        # Run in a separate thread so GUI doesn't freeze
        thread = threading.Thread(target=self.run_generation)
        thread.start()

    # --- LOGIC ---

    def run_generation(self):
        folder = self.folder_path_var.get()
        template_file = self.template_path_var.get()

        if not folder or not template_file:
            messagebox.showerror("Error", "Please select both a folder and a template file.")
            return

        self.generate_btn.config(state='disabled', text="Processing...")

        try:
            # 1. Read Code
            self.log(f"Reading code from: {folder}")
            code_content = self.read_codebase(folder)
            self.log(f"Found {len(code_content)} characters of code.")

            # 2. Read Template
            self.log(f"Reading template: {os.path.basename(template_file)}")
            template_text = self.read_template_file(template_file)

            if not template_text:
                self.log("Error: Could not read text from template.")
                return

            # 3. Send to AI
            self.log("Sending to AI... (Please wait)")
            ai_response = self.generate_with_ai(code_content, template_text)

            # 4. Save Result
            output_filename = "Generated_Docs.docx"
            self.save_to_docx(ai_response, output_filename)

            self.log(f"SUCCESS! Saved to {output_filename}")
            messagebox.showinfo("Success", f"Documentation generated!\nSaved as {output_filename}")

        except Exception as e:
            self.log(f"Error: {str(e)}")
            messagebox.showerror("Error", str(e))

        finally:
            self.generate_btn.config(state='normal', text="GENERATE DOCUMENTATION")

    def read_codebase(self, folder_path):
        code_content = ""
        valid_extensions = ['.py', '.js', '.html', '.css', '.java', '.cpp']
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if any(file.endswith(ext) for ext in valid_extensions):
                    file_path = os.path.join(root, file)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            code_content += f"\n\n--- FILE: {file} ---\n"
                            code_content += f.read()
                    except Exception:
                        pass
        return code_content

    def read_template_file(self, file_path):
        text = ""
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_path.endswith('.pdf'):
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            self.log(f"Template Read Error: {e}")
            return None

    def generate_with_ai(self, code, template):
        prompt = f"""
        You are a technical writer. 
        Task: Write technical documentation for the provided code based on the template structure.

        TEMPLATE STRUCTURE:
        {template}

        CODEBASE:
        {code[:15000]} 
        (Note: Code truncated to 15000 chars to fit limits)

        Output ONLY the documentation content.
        """
        response = openai.ChatCompletion.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content

    def save_to_docx(self, content, filename):
        doc = Document()
        doc.add_heading('AI Generated Documentation', 0)

        # Split by lines and add to doc
        for line in content.split('\n'):
            doc.add_paragraph(line)

        doc.save(filename)


if __name__ == "__main__":
    root = tk.Tk()
    app = DocGeneratorApp(root)
    root.mainloop()